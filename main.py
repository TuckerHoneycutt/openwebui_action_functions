"""
title: Document Style Formatter
author: OpenWebUI Community
version: 1.0.0
description: Extract styling from DOCX/PDF documents and apply it to chat content with a modern GUI interface
required_open_webui_version: 0.5.0
requirements: python-docx>=1.1.0,PyMuPDF>=1.23.0,pdf2docx>=0.5.6,pydantic>=2.0.0
"""

"""
IMPORT DEPENDENCIES
===================
This section handles all required dependencies for the action function.

Required Dependencies:
- python-docx: For DOCX file manipulation
- pdf2docx: For PDF to DOCX conversion (includes PyMuPDF/fitz)
- openwebui: For action decorator (usually pre-installed with OpenWebUI)

Installation:
    pip install python-docx PyMuPDF pdf2docx

The imports are structured to:
1. Import standard library modules first
2. Import third-party modules with error handling
3. Provide clear error messages if dependencies are missing
4. Verify dependencies on module load
"""

import os
import sys
import tempfile
import io
import base64
import json
import uuid
from typing import Dict, List, Any, Optional

# Import Pydantic for Valves class (REQUIRED for OpenWebUI Action structure)
try:
    from pydantic import BaseModel
except ImportError as e:
    error_msg = (
        "ERROR: Missing required dependency 'pydantic'. "
        "Please install it with: pip install pydantic"
    )
    print(error_msg, file=sys.stderr)
    raise ImportError(error_msg) from e

# Import python-docx dependencies (REQUIRED)
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.section import WD_SECTION_START
except ImportError as e:
    error_msg = (
        "ERROR: Missing required dependency 'python-docx'. "
        "Please install it with: pip install python-docx"
    )
    print(error_msg, file=sys.stderr)
    raise ImportError(error_msg) from e

# Import pdf2docx dependency (REQUIRED)
# Note: pdf2docx internally uses PyMuPDF (fitz), so installing pdf2docx
# will also install PyMuPDF as a dependency
try:
    from pdf2docx import Converter
except ImportError as e:
    error_msg = (
        "ERROR: Missing required dependency 'pdf2docx'. "
        "Please install it with: pip install pdf2docx"
    )
    print(error_msg, file=sys.stderr)
    raise ImportError(error_msg) from e


def verify_dependencies() -> Dict[str, bool]:
    """
    Verify that all required dependencies are installed.
    Returns a dictionary with dependency names as keys and installation status as values.
    """
    dependencies = {
        'python-docx': False,
        'pdf2docx': False,
        'openwebui': False
    }

    try:
        import docx
        dependencies['python-docx'] = True
    except ImportError:
        pass

    try:
        import pdf2docx
        dependencies['pdf2docx'] = True
    except ImportError:
        pass

    try:
        import openwebui
        dependencies['openwebui'] = True
    except ImportError:
        pass

    return dependencies


# Verify dependencies on import (but don't fail if optional ones are missing)
_DEPS = verify_dependencies()
if not _DEPS['python-docx']:
    print("WARNING: python-docx is not installed. Install with: pip install python-docx", file=sys.stderr)
if not _DEPS['pdf2docx']:
    print("WARNING: pdf2docx is not installed. Install with: pip install pdf2docx", file=sys.stderr)


def generate_modern_gui() -> str:
    """Generate a modern, stylish GUI HTML with advanced styling and animations."""
    gui_id = f"doc-formatter-{uuid.uuid4().hex[:8]}"

    html = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document Style Formatter</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=Poppins:wght@300;400;500;600;700;800&display=swap');

        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        :root {{
            --primary-gradient: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            --secondary-gradient: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
            --accent-gradient: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
            --dark-gradient: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
            --glass-bg: rgba(255, 255, 255, 0.1);
            --glass-border: rgba(255, 255, 255, 0.2);
            --shadow-lg: 0 20px 60px rgba(0, 0, 0, 0.3);
            --shadow-xl: 0 25px 80px rgba(0, 0, 0, 0.4);
        }}

        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
            background-size: 400% 400%;
            animation: gradientShift 15s ease infinite;
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 20px;
            overflow: hidden;
            position: relative;
        }}

        @keyframes gradientShift {{
            0% {{ background-position: 0% 50%; }}
            50% {{ background-position: 100% 50%; }}
            100% {{ background-position: 0% 50%; }}
        }}

        /* Animated background particles */
        .particles {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            pointer-events: none;
            z-index: 1;
        }}

        .particle {{
            position: absolute;
            width: 4px;
            height: 4px;
            background: rgba(255, 255, 255, 0.5);
            border-radius: 50%;
            animation: float 20s infinite;
        }}

        @keyframes float {{
            0% {{
                transform: translateY(100vh) translateX(0) rotate(0deg);
                opacity: 0;
            }}
            10% {{
                opacity: 1;
            }}
            90% {{
                opacity: 1;
            }}
            100% {{
                transform: translateY(-100vh) translateX(100px) rotate(360deg);
                opacity: 0;
            }}
        }}

        .modal-overlay {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(0, 0, 0, 0.6);
            backdrop-filter: blur(10px);
            z-index: 1000;
            display: flex;
            align-items: center;
            justify-content: center;
            animation: fadeIn 0.3s ease;
        }}

        @keyframes fadeIn {{
            from {{ opacity: 0; }}
            to {{ opacity: 1; }}
        }}

        @keyframes fadeOut {{
            from {{ opacity: 1; }}
            to {{ opacity: 0; }}
        }}

        .modal-container {{
            background: rgba(255, 255, 255, 0.95);
            backdrop-filter: blur(20px);
            border-radius: 24px;
            box-shadow: var(--shadow-xl);
            width: 90%;
            max-width: 600px;
            max-height: 90vh;
            overflow-y: auto;
            position: relative;
            animation: slideUp 0.4s cubic-bezier(0.34, 1.56, 0.64, 1);
            border: 1px solid var(--glass-border);
        }}

        @keyframes slideUp {{
            from {{
                transform: translateY(50px) scale(0.9);
                opacity: 0;
            }}
            to {{
                transform: translateY(0) scale(1);
                opacity: 1;
            }}
        }}

        .modal-header {{
            background: var(--primary-gradient);
            padding: 30px;
            border-radius: 24px 24px 0 0;
            position: relative;
            overflow: hidden;
        }}

        .modal-header::before {{
            content: '';
            position: absolute;
            top: -50%;
            right: -50%;
            width: 200%;
            height: 200%;
            background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%);
            animation: rotate 20s linear infinite;
        }}

        @keyframes rotate {{
            from {{ transform: rotate(0deg); }}
            to {{ transform: rotate(360deg); }}
        }}

        .modal-title {{
            font-family: 'Poppins', sans-serif;
            font-size: 28px;
            font-weight: 700;
            color: white;
            margin-bottom: 8px;
            position: relative;
            z-index: 1;
            text-shadow: 0 2px 10px rgba(0, 0, 0, 0.2);
            background: linear-gradient(135deg, #ffffff 0%, #f0f0f0 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            animation: titleGlow 3s ease-in-out infinite;
        }}

        @keyframes titleGlow {{
            0%, 100% {{ filter: brightness(1); }}
            50% {{ filter: brightness(1.2); }}
        }}

        .modal-subtitle {{
            font-size: 14px;
            color: rgba(255, 255, 255, 0.9);
            position: relative;
            z-index: 1;
            font-weight: 400;
        }}

        .close-btn {{
            position: absolute;
            top: 20px;
            right: 20px;
            width: 40px;
            height: 40px;
            border-radius: 50%;
            background: rgba(255, 255, 255, 0.2);
            border: none;
            color: white;
            font-size: 24px;
            cursor: pointer;
            display: flex;
            align-items: center;
            justify-content: center;
            transition: all 0.3s ease;
            z-index: 2;
        }}

        .close-btn:hover {{
            background: rgba(255, 255, 255, 0.3);
            transform: rotate(90deg) scale(1.1);
        }}

        .modal-body {{
            padding: 40px;
        }}

        .upload-area {{
            border: 3px dashed #667eea;
            border-radius: 16px;
            padding: 60px 40px;
            text-align: center;
            background: linear-gradient(135deg, rgba(102, 126, 234, 0.05) 0%, rgba(118, 75, 162, 0.05) 100%);
            transition: all 0.3s ease;
            cursor: pointer;
            position: relative;
            overflow: hidden;
        }}

        .upload-area::before {{
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
            transition: left 0.5s;
        }}

        .upload-area:hover::before {{
            left: 100%;
        }}

        .upload-area:hover {{
            border-color: #764ba2;
            background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
            transform: translateY(-2px);
            box-shadow: 0 10px 30px rgba(102, 126, 234, 0.2);
        }}

        .upload-area.dragover {{
            border-color: #4facfe;
            background: linear-gradient(135deg, rgba(79, 172, 254, 0.15) 0%, rgba(0, 242, 254, 0.15) 100%);
            transform: scale(1.02);
        }}

        .upload-icon {{
            width: 80px;
            height: 80px;
            margin: 0 auto 20px;
            background: var(--primary-gradient);
            border-radius: 20px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 40px;
            box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
            animation: pulse 2s ease-in-out infinite;
            position: relative;
            overflow: hidden;
        }}

        .upload-icon::before {{
            content: '';
            position: absolute;
            top: -50%;
            left: -50%;
            width: 200%;
            height: 200%;
            background: linear-gradient(45deg, transparent, rgba(255,255,255,0.3), transparent);
            animation: shine 3s infinite;
        }}

        @keyframes shine {{
            0% {{ transform: translateX(-100%) translateY(-100%) rotate(45deg); }}
            100% {{ transform: translateX(100%) translateY(100%) rotate(45deg); }}
        }}

        @keyframes pulse {{
            0%, 100% {{ transform: scale(1); }}
            50% {{ transform: scale(1.05); }}
        }}

        .upload-text {{
            font-size: 18px;
            font-weight: 600;
            color: #333;
            margin-bottom: 8px;
        }}

        .upload-hint {{
            font-size: 14px;
            color: #666;
            margin-top: 8px;
        }}

        .file-input {{
            display: none;
        }}

        .file-info {{
            margin-top: 20px;
            padding: 20px;
            background: linear-gradient(135deg, rgba(79, 172, 254, 0.1) 0%, rgba(0, 242, 254, 0.1) 100%);
            border-radius: 12px;
            display: none;
            animation: slideDown 0.3s ease;
        }}

        @keyframes slideDown {{
            from {{
                opacity: 0;
                transform: translateY(-10px);
            }}
            to {{
                opacity: 1;
                transform: translateY(0);
            }}
        }}

        .file-info.show {{
            display: block;
        }}

        .file-name {{
            font-weight: 600;
            color: #333;
            margin-bottom: 4px;
            display: flex;
            align-items: center;
            gap: 10px;
        }}

        .file-size {{
            font-size: 12px;
            color: #666;
        }}

        .progress-container {{
            margin-top: 30px;
            display: none;
        }}

        .progress-container.show {{
            display: block;
        }}

        .progress-label {{
            font-size: 14px;
            font-weight: 600;
            color: #333;
            margin-bottom: 10px;
            display: flex;
            justify-content: space-between;
        }}

        .progress-bar {{
            width: 100%;
            height: 12px;
            background: rgba(102, 126, 234, 0.1);
            border-radius: 10px;
            overflow: hidden;
            position: relative;
        }}

        .progress-fill {{
            height: 100%;
            background: var(--accent-gradient);
            border-radius: 10px;
            width: 0%;
            transition: width 0.3s ease;
            position: relative;
            overflow: hidden;
        }}

        .progress-fill::after {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            bottom: 0;
            right: 0;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.3), transparent);
            animation: shimmer 2s infinite;
        }}

        @keyframes shimmer {{
            0% {{ transform: translateX(-100%); }}
            100% {{ transform: translateX(100%); }}
        }}

        .action-buttons {{
            display: flex;
            gap: 15px;
            margin-top: 30px;
        }}

        .btn {{
            flex: 1;
            padding: 16px 32px;
            border: none;
            border-radius: 12px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
            position: relative;
            overflow: hidden;
            font-family: 'Inter', sans-serif;
        }}

        .btn::before {{
            content: '';
            position: absolute;
            top: 50%;
            left: 50%;
            width: 0;
            height: 0;
            border-radius: 50%;
            background: rgba(255, 255, 255, 0.3);
            transform: translate(-50%, -50%);
            transition: width 0.6s, height 0.6s;
        }}

        .btn:hover::before {{
            width: 300px;
            height: 300px;
        }}

        .btn-primary {{
            background: var(--primary-gradient);
            color: white;
            box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
        }}

        .btn-primary:hover {{
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(102, 126, 234, 0.5);
        }}

        .btn-primary:active {{
            transform: translateY(0);
        }}

        .btn-secondary {{
            background: rgba(102, 126, 234, 0.1);
            color: #667eea;
            border: 2px solid #667eea;
        }}

        .btn-secondary:hover {{
            background: rgba(102, 126, 234, 0.2);
        }}

        .btn:disabled {{
            opacity: 0.6;
            cursor: not-allowed;
        }}

        .status-message {{
            margin-top: 20px;
            padding: 16px;
            border-radius: 12px;
            display: none;
            animation: slideDown 0.3s ease;
        }}

        .status-message.show {{
            display: block;
        }}

        .status-success {{
            background: linear-gradient(135deg, rgba(76, 175, 80, 0.1) 0%, rgba(56, 142, 60, 0.1) 100%);
            border: 2px solid #4caf50;
            color: #2e7d32;
        }}

        .status-error {{
            background: linear-gradient(135deg, rgba(244, 67, 54, 0.1) 0%, rgba(198, 40, 40, 0.1) 100%);
            border: 2px solid #f44336;
            color: #c62828;
        }}

        .status-info {{
            background: linear-gradient(135deg, rgba(33, 150, 243, 0.1) 0%, rgba(25, 118, 210, 0.1) 100%);
            border: 2px solid #2196f3;
            color: #1976d2;
        }}

        .spinner {{
            display: inline-block;
            width: 20px;
            height: 20px;
            border: 3px solid rgba(255, 255, 255, 0.3);
            border-top-color: white;
            border-radius: 50%;
            animation: spin 0.8s linear infinite;
            margin-right: 10px;
        }}

        @keyframes spin {{
            to {{ transform: rotate(360deg); }}
        }}

        /* Responsive design */
        @media (max-width: 768px) {{
            .modal-container {{
                width: 95%;
                max-height: 95vh;
            }}

            .modal-header {{
                padding: 20px;
            }}

            .modal-title {{
                font-size: 22px;
            }}

            .modal-body {{
                padding: 25px;
            }}

            .upload-area {{
                padding: 40px 20px;
            }}

            .action-buttons {{
                flex-direction: column;
            }}
        }}

        /* Scrollbar styling */
        .modal-container::-webkit-scrollbar {{
            width: 8px;
        }}

        .modal-container::-webkit-scrollbar-track {{
            background: rgba(0, 0, 0, 0.05);
            border-radius: 10px;
        }}

        .modal-container::-webkit-scrollbar-thumb {{
            background: var(--primary-gradient);
            border-radius: 10px;
        }}

        .modal-container::-webkit-scrollbar-thumb:hover {{
            background: var(--secondary-gradient);
        }}
    </style>
</head>
<body>
    <div class="particles" id="particles"></div>

    <div class="modal-overlay" id="modalOverlay-{gui_id}">
        <div class="modal-container">
            <div class="modal-header">
                <button class="close-btn" onclick="window.closeDocFormatterModal_{gui_id}()">×</button>
                <h1 class="modal-title">📄 Document Style Formatter</h1>
                <p class="modal-subtitle">Transform your chat into a professionally styled document</p>
            </div>

            <div class="modal-body">
                <div class="upload-area" id="uploadArea" onclick="document.getElementById('fileInput').click()">
                    <div class="upload-icon">📎</div>
                    <div class="upload-text">Drag & Drop Your Document</div>
                    <div class="upload-hint">or click to browse</div>
                    <div class="upload-hint" style="margin-top: 10px; font-size: 12px; color: #999;">
                        Supports DOCX and PDF files
                    </div>
                </div>

                <input type="file" id="fileInput" class="file-input" accept=".docx,.pdf" />

                <div class="file-info" id="fileInfo">
                    <div class="file-name">
                        <span>📄</span>
                        <span id="fileName"></span>
                    </div>
                    <div class="file-size" id="fileSize"></span>
                </div>

                <div class="progress-container" id="progressContainer">
                    <div class="progress-label">
                        <span>Processing...</span>
                        <span id="progressPercent">0%</span>
                    </div>
                    <div class="progress-bar">
                        <div class="progress-fill" id="progressFill"></div>
                    </div>
                </div>

                <div class="status-message" id="statusMessage"></div>

                <div class="action-buttons">
                    <button class="btn btn-primary" id="processBtn" onclick="processDocument()" disabled>
                        <span id="processBtnText">Process Document</span>
                    </button>
                    <button class="btn btn-secondary" onclick="window.closeDocFormatterModal_{gui_id}()">Cancel</button>
                </div>
            </div>
        </div>
    </div>

    <script>
        const guiId = '{gui_id}';
        let selectedFile = null;

        // Try to get chat messages from OpenWebUI context
        async function getChatMessages() {{
            try {{
                // Method 1: Check if messages were passed in the response data
                const responseData = window.__DOC_FORMATTER_CHAT_MESSAGES__;
                if (responseData && Array.isArray(responseData)) {{
                    return responseData;
                }}

                // Method 2: From window context
                if (window.chatMessages) {{
                    return window.chatMessages;
                }}

                // Method 3: From parent window
                if (window.parent && window.parent.chatMessages) {{
                    return window.parent.chatMessages;
                }}

                // Method 4: Try to access OpenWebUI's chat state
                if (window.__OPENWEBUI_CHAT_MESSAGES__) {{
                    return window.__OPENWEBUI_CHAT_MESSAGES__;
                }}

                // Method 5: Try to find messages in DOM (look for OpenWebUI message elements)
                const chatElements = document.querySelectorAll('[data-message], .message, .chat-message, [class*="message"]');
                if (chatElements.length > 0) {{
                    const messages = [];
                    chatElements.forEach(el => {{
                        const role = el.getAttribute('data-role') ||
                                   el.getAttribute('data-from') ||
                                   (el.classList.contains('user') || el.querySelector('.user')) ? 'user' :
                                   (el.classList.contains('assistant') || el.querySelector('.assistant')) ? 'assistant' :
                                   'user';
                        const contentEl = el.querySelector('.message-content, .content, [class*="content"]') || el;
                        const content = contentEl.textContent || contentEl.innerText || '';
                        if (content.trim()) {{
                            messages.push({{ role: role, content: content.trim() }});
                        }}
                    }});
                    if (messages.length > 0) {{
                        return messages;
                    }}
                }}

                // Method 6: Try to get from localStorage or sessionStorage
                try {{
                    const stored = localStorage.getItem('openwebui_chat_messages') ||
                                 sessionStorage.getItem('openwebui_chat_messages');
                    if (stored) {{
                        return JSON.parse(stored);
                    }}
                }} catch (e) {{
                    console.log('Could not parse stored messages:', e);
                }}

                console.warn('Could not find chat messages. The formatted document may be empty.');
                return [];
            }} catch (e) {{
                console.log('Could not access chat messages:', e);
                return [];
            }}
        }}

        // Store function globally for use in processDocument
        window.getChatMessages = getChatMessages;

        // Create animated particles
        function createParticles() {{
            const particlesContainer = document.getElementById('particles');
            const particleCount = 30;

            for (let i = 0; i < particleCount; i++) {{
                const particle = document.createElement('div');
                particle.className = 'particle';
                particle.style.left = Math.random() * 100 + '%';
                particle.style.animationDelay = Math.random() * 20 + 's';
                particle.style.animationDuration = (15 + Math.random() * 10) + 's';
                particlesContainer.appendChild(particle);
            }}
        }}

        // File input handling
        const fileInput = document.getElementById('fileInput');
        const uploadArea = document.getElementById('uploadArea');
        const fileInfo = document.getElementById('fileInfo');
        const processBtn = document.getElementById('processBtn');

        fileInput.addEventListener('change', handleFileSelect);

        uploadArea.addEventListener('dragover', (e) => {{
            e.preventDefault();
            uploadArea.classList.add('dragover');
        }});

        uploadArea.addEventListener('dragleave', () => {{
            uploadArea.classList.remove('dragover');
        }});

        uploadArea.addEventListener('drop', (e) => {{
            e.preventDefault();
            uploadArea.classList.remove('dragover');
            const files = e.dataTransfer.files;
            if (files.length > 0) {{
                handleFile(files[0]);
            }}
        }});

        function handleFileSelect(e) {{
            if (e.target.files.length > 0) {{
                handleFile(e.target.files[0]);
            }}
        }}

        function handleFile(file) {{
            const validTypes = ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/pdf'];
            const validExtensions = ['.docx', '.pdf'];
            const fileExt = '.' + file.name.split('.').pop().toLowerCase();

            if (!validExtensions.includes(fileExt) && !validTypes.includes(file.type)) {{
                showStatus('Please select a DOCX or PDF file.', 'error');
                return;
            }}

            selectedFile = file;
            document.getElementById('fileName').textContent = file.name;
            document.getElementById('fileSize').textContent = formatFileSize(file.size);
            fileInfo.classList.add('show');
            processBtn.disabled = false;
        }}

        function formatFileSize(bytes) {{
            if (bytes === 0) return '0 Bytes';
            const k = 1024;
            const sizes = ['Bytes', 'KB', 'MB', 'GB'];
            const i = Math.floor(Math.log(bytes) / Math.log(k));
            return Math.round(bytes / Math.pow(k, i) * 100) / 100 + ' ' + sizes[i];
        }}

        function updateProgress(percent) {{
            const progressFill = document.getElementById('progressFill');
            const progressPercent = document.getElementById('progressPercent');
            const progressContainer = document.getElementById('progressContainer');

            progressFill.style.width = percent + '%';
            progressPercent.textContent = Math.round(percent) + '%';
            progressContainer.classList.add('show');
        }}

        function showStatus(message, type) {{
            const statusMessage = document.getElementById('statusMessage');
            statusMessage.textContent = message;
            statusMessage.className = 'status-message show status-' + type;
        }}

        async function processDocument() {{
            if (!selectedFile) {{
                showStatus('Please select a file first.', 'error');
                return;
            }}

            processBtn.disabled = true;
            const processBtnText = document.getElementById('processBtnText');
            processBtnText.innerHTML = '<span class="spinner"></span>Processing...';

            updateProgress(10);
            showStatus('Reading file...', 'info');

            try {{
                const formData = new FormData();
                formData.append('file', selectedFile);

                // Simulate progress
                setTimeout(() => updateProgress(30), 500);
                setTimeout(() => updateProgress(50), 1000);
                setTimeout(() => updateProgress(70), 1500);

                // Get chat messages
                updateProgress(20);
                const chatMessages = await getChatMessages();
                updateProgress(25);

                // Convert file to base64 for transmission
                const fileBase64 = await fileToBase64(selectedFile);
                updateProgress(35);

                // Try to call the action function through OpenWebUI's API
                // OpenWebUI typically exposes actions through a specific endpoint
                const actionName = 'format_chat_with_document_style';
                const apiEndpoints = [
                    `/api/v1/actions/${{actionName}}`,
                    `/api/actions/${{actionName}}`,
                    `/api/functions/${{actionName}}`,
                    window.location.origin + `/api/v1/actions/${{actionName}}`
                ];

                let result = null;
                let errorOccurred = false;
                let lastError = null;

                for (const endpoint of apiEndpoints) {{
                    try {{
                        updateProgress(45);
                        const response = await fetch(endpoint, {{
                            method: 'POST',
                            headers: {{
                                'Content-Type': 'application/json',
                            }},
                            body: JSON.stringify({{
                                file: fileBase64,
                                file_extension: '.' + selectedFile.name.split('.').pop().toLowerCase(),
                                messages: chatMessages,
                                chat_messages: chatMessages
                            }})
                        }});

                        updateProgress(65);

                        if (response.ok) {{
                            result = await response.json();
                            break;
                        }} else {{
                            const errorData = await response.json().catch(() => ({{}}));
                            lastError = errorData.error || `HTTP ${{response.status}}: ${{response.statusText}}`;
                            if (endpoint === apiEndpoints[apiEndpoints.length - 1]) {{
                                throw new Error(lastError);
                            }}
                        }}
                    }} catch (err) {{
                        lastError = err.message;
                        if (endpoint === apiEndpoints[apiEndpoints.length - 1]) {{
                            errorOccurred = true;
                            // If all API calls fail, try calling the function directly via window
                            try {{
                                if (window.__OPENWEBUI_CALL_ACTION__) {{
                                    updateProgress(50);
                                    result = await window.__OPENWEBUI_CALL_ACTION__(actionName, {{
                                        file: fileBase64,
                                        file_extension: '.' + selectedFile.name.split('.').pop().toLowerCase(),
                                        messages: chatMessages,
                                        chat_messages: chatMessages
                                    }});
                                    updateProgress(80);
                                }} else {{
                                    throw err;
                                }}
                            }} catch (fallbackErr) {{
                                throw err;
                            }}
                        }}
                    }}
                }}

                updateProgress(90);

                if (result && !errorOccurred) {{
                    updateProgress(100);

                    if (result.error) {{
                        showStatus(result.error, 'error');
                        processBtn.disabled = false;
                        processBtnText.textContent = 'Process Document';
                        return;
                    }}

                    showStatus('Document formatted successfully! Download starting...', 'success');

                    // Trigger download
                    if (result.file && result.file.content) {{
                        const blob = base64ToBlob(result.file.content, result.file.mime_type || 'application/vnd.openxmlformats-officedocument.wordprocessingml.document');
                        const url = URL.createObjectURL(blob);
                        const a = document.createElement('a');
                        a.href = url;
                        a.download = result.file.filename || 'formatted_chat.docx';
                        document.body.appendChild(a);
                        a.click();
                        document.body.removeChild(a);
                        URL.revokeObjectURL(url);
                    }} else if (result.download_url) {{
                        window.open(result.download_url, '_blank');
                    }}

                    setTimeout(() => {{
                        closeModal();
                    }}, 2000);
                }}
            }} catch (error) {{
                showStatus('Error: ' + error.message, 'error');
                processBtn.disabled = false;
                processBtnText.textContent = 'Process Document';
            }}
        }}

        async function fileToBase64(file) {{
            return new Promise((resolve, reject) => {{
                const reader = new FileReader();
                reader.onload = () => {{
                    const base64 = reader.result.split(',')[1];
                    resolve(base64);
                }};
                reader.onerror = reject;
                reader.readAsDataURL(file);
            }});
        }}

        function base64ToBlob(base64, mimeType) {{
            const byteCharacters = atob(base64);
            const byteNumbers = new Array(byteCharacters.length);
            for (let i = 0; i < byteCharacters.length; i++) {{
                byteNumbers[i] = byteCharacters.charCodeAt(i);
            }}
            const byteArray = new Uint8Array(byteNumbers);
            return new Blob([byteArray], {{ type: mimeType }});
        }}

        function closeModal_{gui_id}() {{
            const overlay = document.getElementById('modalOverlay-{gui_id}');
            if (overlay) {{
                overlay.style.animation = 'fadeOut 0.3s ease';
                setTimeout(() => {{
                    overlay.remove();
                }}, 300);
            }}
        }}

        // Make function globally accessible
        window.closeDocFormatterModal_{gui_id} = closeModal_{gui_id};

        // Also create alias for backward compatibility
        function closeModal() {{
            closeModal_{gui_id}();
        }}

        // Store chat messages if available from response
        try {{
            // Try to get messages from parent response or context
            const responseElement = document.querySelector('[data-chat-messages], [data-messages]');
            if (responseElement) {{
                const messagesData = responseElement.getAttribute('data-chat-messages') ||
                                   responseElement.getAttribute('data-messages');
                if (messagesData) {{
                    window.__DOC_FORMATTER_CHAT_MESSAGES__ = JSON.parse(messagesData);
                }}
            }}
        }} catch (e) {{
            console.log('Could not parse messages from DOM:', e);
        }}

        // Initialize
        createParticles();

        // Close on overlay click
        const overlayEl = document.getElementById('modalOverlay-{gui_id}');
        if (overlayEl) {{
            overlayEl.addEventListener('click', (e) => {{
                if (e.target.id === 'modalOverlay-{gui_id}') {{
                    closeModal_{gui_id}();
                }}
            }});
        }}

        // Store chat messages if passed in response
        if (typeof window.__DOC_FORMATTER_RESPONSE__ !== 'undefined') {{
            const response = window.__DOC_FORMATTER_RESPONSE__;
            if (response._chat_messages) {{
                window.__DOC_FORMATTER_CHAT_MESSAGES__ = response._chat_messages;
            }}
        }}
    </script>
</body>
</html>
"""
    return html


class DocumentStyleExtractor:
    """Extracts styling information from DOCX and PDF documents."""

    def __init__(self):
        self.styles = {
            'default_font': {'name': 'Calibri', 'size': Pt(11), 'color': RGBColor(0, 0, 0), 'bold': False, 'italic': False, 'underline': False},
            'heading_styles': [],
            'paragraph_styles': [],
            'headers': [],
            'footers': [],
            'tables': [],
            'page_breaks': [],
            'sections': [],
            'margins': {'top': None, 'bottom': None, 'left': None, 'right': None},
            'indentations': {'left': None, 'right': None, 'first_line': None, 'hanging': None},
            'line_spacing': None,
            'alignment': WD_PARAGRAPH_ALIGNMENT.LEFT
        }

    def extract_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """Extract styles from a DOCX file."""
        doc = Document(docx_path)

        # Extract section properties (margins, page breaks)
        for i, section in enumerate(doc.sections):
            section_info = {
                'margins': {
                    'top': section.top_margin,
                    'bottom': section.bottom_margin,
                    'left': section.left_margin,
                    'right': section.right_margin
                },
                'page_break_before': section.start_type == WD_SECTION_START.NEW_PAGE
            }
            self.styles['sections'].append(section_info)

            if i == 0:  # Use first section's margins as default
                self.styles['margins'] = section_info['margins']

        # Extract headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    header_style = self._extract_paragraph_style(paragraph)
                    self.styles['headers'].append({
                        'text': paragraph.text,
                        'style': header_style
                    })

        # Extract footers
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    footer_style = self._extract_paragraph_style(paragraph)
                    self.styles['footers'].append({
                        'text': paragraph.text,
                        'style': footer_style
                    })

        # Extract paragraph styles
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                para_style = self._extract_paragraph_style(paragraph)

                # Check if it's a heading
                if paragraph.style.name.startswith('Heading'):
                    self.styles['heading_styles'].append(para_style)
                else:
                    self.styles['paragraph_styles'].append(para_style)

        # Extract table styles
        for table in doc.tables:
            table_data = {
                'rows': [],
                'style': None
            }
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text
                    cell_style = self._extract_paragraph_style(cell.paragraphs[0]) if cell.paragraphs else None
                    row_data.append({
                        'text': cell_text,
                        'style': cell_style
                    })
                table_data['rows'].append(row_data)
            self.styles['tables'].append(table_data)

        # Set defaults from first paragraph if available
        if doc.paragraphs:
            first_para = doc.paragraphs[0]
            if first_para.runs:
                run = first_para.runs[0]
                if run.font.name:
                    self.styles['default_font']['name'] = run.font.name
                if run.font.size:
                    self.styles['default_font']['size'] = run.font.size
                if run.font.color and run.font.color.rgb:
                    self.styles['default_font']['color'] = run.font.color.rgb
                if run.font.bold is not None:
                    self.styles['default_font']['bold'] = run.font.bold
                if run.font.italic is not None:
                    self.styles['default_font']['italic'] = run.font.italic
                if run.font.underline is not None:
                    self.styles['default_font']['underline'] = run.font.underline

            if first_para.paragraph_format.line_spacing:
                self.styles['line_spacing'] = first_para.paragraph_format.line_spacing
            if first_para.alignment:
                self.styles['alignment'] = first_para.alignment
            if first_para.paragraph_format.left_indent:
                self.styles['indentations']['left'] = first_para.paragraph_format.left_indent
            if first_para.paragraph_format.right_indent:
                self.styles['indentations']['right'] = first_para.paragraph_format.right_indent
            if first_para.paragraph_format.first_line_indent:
                self.styles['indentations']['first_line'] = first_para.paragraph_format.first_line_indent

        return self.styles


class DocumentStyleApplier:
    """Applies extracted styles to chat content in a new document."""

    def __init__(self, styles: Dict[str, Any]):
        self.styles = styles
        self.doc = Document()
        self._apply_document_settings()

    def _extract_paragraph_style(self, paragraph) -> Dict[str, Any]:
        """Extract style information from a paragraph."""
        style_info = {
            'font': {},
            'alignment': paragraph.alignment,
            'line_spacing': paragraph.paragraph_format.line_spacing,
            'indentation': {
                'left': paragraph.paragraph_format.left_indent,
                'right': paragraph.paragraph_format.right_indent,
                'first_line': paragraph.paragraph_format.first_line_indent,
                'hanging': paragraph.paragraph_format.hanging_indent
            },
            'space_before': paragraph.paragraph_format.space_before,
            'space_after': paragraph.paragraph_format.space_after
        }

        # Extract font information from runs
        if paragraph.runs:
            first_run = paragraph.runs[0]
            style_info['font'] = {
                'name': first_run.font.name,
                'size': first_run.font.size,
                'color': first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None,
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline
            }

        return style_info


class DocumentStyleApplier:
    """Applies extracted styles to chat content in a new document."""

    def __init__(self, styles: Dict[str, Any]):
        self.styles = styles
        self.doc = Document()
        self._apply_document_settings()

    def _apply_document_settings(self):
        """Apply document-level settings like margins."""
        if self.styles['sections']:
            section = self.doc.sections[0]
            margins = self.styles['margins']
            if margins['top']:
                section.top_margin = margins['top']
            if margins['bottom']:
                section.bottom_margin = margins['bottom']
            if margins['left']:
                section.left_margin = margins['left']
            if margins['right']:
                section.right_margin = margins['right']

    def add_headers(self):
        """Add headers from the source document."""
        for header_info in self.styles['headers']:
            header_para = self.doc.sections[0].header.paragraphs[0] if self.doc.sections[0].header.paragraphs else self.doc.sections[0].header.add_paragraph()
            header_para.text = header_info['text']
            self._apply_style_to_paragraph(header_para, header_info['style'])

    def add_footers(self):
        """Add footers from the source document."""
        for footer_info in self.styles['footers']:
            footer_para = self.doc.sections[0].footer.paragraphs[0] if self.doc.sections[0].footer.paragraphs else self.doc.sections[0].footer.add_paragraph()
            footer_para.text = footer_info['text']
            self._apply_style_to_paragraph(footer_para, footer_info['style'])

    def add_chat_content(self, chat_messages: List[Dict[str, str]]):
        """Add chat content with applied styles."""
        default_font = self.styles['default_font']
        para_style = self.styles['paragraph_styles'][0] if self.styles['paragraph_styles'] else None

        for i, message in enumerate(chat_messages):
            role = message.get('role', 'user')
            content = message.get('content', '')

            # Add role label
            role_para = self.doc.add_paragraph()
            role_run = role_para.add_run(f"{role.capitalize()}: ")
            self._apply_font_to_run(role_run, default_font)
            role_run.bold = True

            # Add content
            content_para = self.doc.add_paragraph()
            content_run = content_para.add_run(content)

            # Apply paragraph style if available
            if para_style:
                self._apply_style_to_paragraph(content_para, para_style)
            else:
                self._apply_font_to_run(content_run, default_font)

            # Add spacing between messages
            if i < len(chat_messages) - 1:
                self.doc.add_paragraph()

    def add_tables(self):
        """Add tables from the source document."""
        for table_info in self.styles['tables']:
            if not table_info['rows']:
                continue

            num_cols = len(table_info['rows'][0])
            table = self.doc.add_table(rows=0, cols=num_cols)

            for row_data in table_info['rows']:
                row = table.add_row()
                for j, cell_data in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = cell_data['text']
                    if cell_data['style'] and cell.paragraphs:
                        self._apply_style_to_paragraph(cell.paragraphs[0], cell_data['style'])

            self.doc.add_paragraph()  # Add spacing after table

    def _apply_style_to_paragraph(self, paragraph, style_info: Dict[str, Any]):
        """Apply style information to a paragraph."""
        if not style_info:
            return

        # Apply alignment
        if style_info.get('alignment'):
            paragraph.alignment = style_info['alignment']

        # Apply line spacing
        if style_info.get('line_spacing'):
            paragraph.paragraph_format.line_spacing = style_info['line_spacing']

        # Apply indentation
        indent = style_info.get('indentation', {})
        if indent.get('left'):
            paragraph.paragraph_format.left_indent = indent['left']
        if indent.get('right'):
            paragraph.paragraph_format.right_indent = indent['right']
        if indent.get('first_line'):
            paragraph.paragraph_format.first_line_indent = indent['first_line']
        if indent.get('hanging'):
            paragraph.paragraph_format.hanging_indent = indent['hanging']

        # Apply spacing
        if style_info.get('space_before'):
            paragraph.paragraph_format.space_before = style_info['space_before']
        if style_info.get('space_after'):
            paragraph.paragraph_format.space_after = style_info['space_after']

        # Apply font to runs
        font_info = style_info.get('font', {})
        if paragraph.runs:
            for run in paragraph.runs:
                self._apply_font_to_run(run, font_info)

    def _apply_font_to_run(self, run, font_info: Dict[str, Any]):
        """Apply font information to a run."""
        if not font_info:
            return

        if font_info.get('name'):
            run.font.name = font_info['name']
        if font_info.get('size'):
            run.font.size = font_info['size']
        if font_info.get('color'):
            run.font.color.rgb = font_info['color']
        if font_info.get('bold') is not None:
            run.font.bold = font_info['bold']
        if font_info.get('italic') is not None:
            run.font.italic = font_info['italic']
        if font_info.get('underline') is not None:
            run.font.underline = font_info['underline']

    def save(self, output_path: str):
        """Save the document."""
        self.doc.save(output_path)

    def get_bytes(self) -> bytes:
        """Get document as bytes."""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()


class Action:
    """
    OpenWebUI Action class for Document Style Formatter.
    Extracts styling from DOCX/PDF documents and applies it to chat content.
    """

    class Valves(BaseModel):
        """Configuration parameters for the action."""
        pass

    def __init__(self):
        """Initialize the Action with Valves configuration."""
        self.valves = self.Valves()

    def _get_file_upload_modal(self, chat_msgs: List[Dict[str, Any]]) -> str:
        """Generate JavaScript code to create a file upload modal (like prior.py approach)"""
        import json
        chat_msgs_json = json.dumps(chat_msgs)

        return f"""
            // Store chat messages globally for use when processing
            window.docFormatterChatMsgs = {chat_msgs_json};
            const overlay = document.createElement('div');
            overlay.style.cssText = `position: fixed; top: 0; left: 0; width: 100%; height: 100%; background-color: rgba(0,0,0,0.5); display: flex; justify-content: center; align-items: center; z-index: 10000;`;
            const modal = document.createElement('div');
            modal.style.cssText = `background: white; padding: 30px; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.3); max-width: 500px; width: 90%; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-height: 90vh; overflow-y: auto;`;

            modal.innerHTML = `
                <h2 style="margin-top: 0; color: #667eea; text-align: center; font-size: 24px; font-weight: 600;">📄 Document Style Formatter</h2>
                <p style="color: #666; margin-bottom: 20px; text-align: center;">Upload a DOCX or PDF document to extract styling and format your chat content.</p>

                <div style="margin-bottom: 20px;">
                    <label style="display: block; margin-bottom: 8px; color: #333; font-weight: 600;">Select Document:</label>
                    <input type="file" id="docFormatterFileInput" accept=".docx,.pdf" style="width: 100%; padding: 12px; font-size: 14px; border: 2px solid #667eea; border-radius: 8px; margin-bottom: 10px; cursor: pointer;">
                </div>

                <div id="docFormatterStatus" style="margin: 15px 0; padding: 12px; border-radius: 8px; display: none;"></div>

                <div style="display: flex; justify-content: space-between; gap: 10px;">
                    <button id="processBtn" style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; border: none; padding: 12px 24px; border-radius: 8px; cursor: pointer; flex-grow: 1; font-size: 16px; font-weight: 600;">Process Document</button>
                    <button id="cancelBtn" style="background: #f5f5f5; border: 1px solid #ddd; padding: 12px 24px; border-radius: 8px; cursor: pointer; flex-grow: 1; font-size: 16px;">Cancel</button>
                </div>
            `;

            overlay.appendChild(modal);
            document.body.appendChild(overlay);

            // Store chat messages and model for later use
            window.docFormatterChatMsgs = {chat_msgs_json};
            // Try to get current model from page
            try {{
                const modelSelect = document.querySelector('select[name="model"], [data-model]');
                if (modelSelect) {{
                    window.docFormatterModel = modelSelect.value || modelSelect.getAttribute('data-model');
                }}
            }} catch (e) {{
                window.docFormatterModel = null;
            }}

            const closeModal = () => {{
                if (document.body.contains(overlay)) {{
                    document.body.removeChild(overlay);
                }}
            }};

            const showStatus = (message, isError = false) => {{
                const statusDiv = document.getElementById('docFormatterStatus');
                if (!statusDiv) {{
                    // Element doesn't exist (modal might be closed), just log to console
                    console.log(isError ? '[ERROR]' : '[INFO]', message);
                    return;
                }}
                statusDiv.style.display = 'block';
                statusDiv.style.background = isError ? '#fee' : '#efe';
                statusDiv.style.color = isError ? '#c33' : '#3c3';
                statusDiv.style.border = `1px solid ${{isError ? '#c33' : '#3c3'}}`;
                statusDiv.textContent = message;
                setTimeout(() => {{
                    if (statusDiv && document.body.contains(statusDiv)) {{
                        statusDiv.style.display = 'none';
                    }}
                }}, 5000);
            }};

            // Handle file upload and processing
            document.getElementById('processBtn').onclick = async () => {{
                const fileInput = document.getElementById('docFormatterFileInput');
                const statusDiv = document.getElementById('docFormatterStatus');

                if (!fileInput.files || !fileInput.files[0]) {{
                    showStatus('Please select a file first.', true);
                    return;
                }}

                const file = fileInput.files[0];
                showStatus('Processing file: ' + file.name + '...', false);

                // Convert file to base64
                const reader = new FileReader();
                reader.onload = async function(e) {{
                    const base64 = e.target.result.split(',')[1];

                    try {{
                        showStatus('Processing file...', false);

                        // Prepare file data
                        const fileData = {{
                            file: base64,
                            file_extension: '.' + file.name.split('.').pop().toLowerCase(),
                            messages: window.docFormatterChatMsgs || []
                        }};

                        // Store file data globally
                        window.docFormatterFileData = fileData;
                        localStorage.setItem('docFormatterFileData', JSON.stringify(fileData));

                        // Show status before closing modal
                        showStatus('File data prepared. Processing...', false);

                        // Close modal after a brief delay to show status
                        setTimeout(() => {{
                            closeModal();

                            // Trigger the action function call with file data after modal closes
                            console.log('Preparing to trigger action function with file data');

                            // Create a custom event to trigger action processing
                            // This simulates clicking the action button but with file data included
                            const actionEvent = new CustomEvent('openwebui:action:call', {{
                                detail: {{
                                    action: 'output_to_document',
                                    data: fileData
                                }}
                            }});
                            document.dispatchEvent(actionEvent);

                            // Also try to find and trigger the action button programmatically
                            // Look for action buttons and simulate a click
                            setTimeout(() => {{
                                // Try multiple ways to find the action button
                                const actionSelectors = [
                                    '[data-action="output_to_document"]',
                                    'button[onclick*="output_to_document"]',
                                    '.action-button',
                                    '[data-function="output_to_document"]'
                                ];

                                let actionButton = null;
                                for (const selector of actionSelectors) {{
                                    actionButton = document.querySelector(selector);
                                    if (actionButton) break;
                                }}

                                if (actionButton) {{
                                    console.log('Found action button, triggering click with file data');
                                    // Store file data so it's available when button is clicked
                                    window.docFormatterPendingProcess = true;
                                    actionButton.click();
                                }} else {{
                                    console.log('Action button not found, trying alternative method');
                                    // Alternative: Use fetch interceptor and show notification
                                    const notification = document.createElement('div');
                                    notification.style.cssText = 'position: fixed; top: 20px; right: 20px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 15px 25px; border-radius: 8px; z-index: 10001; box-shadow: 0 2px 10px rgba(0,0,0,0.3); font-weight: 600;';
                                    notification.textContent = '✓ File ready! Click the action button to process.';
                                    document.body.appendChild(notification);

                                    // Install fetch interceptor
                                    const originalFetch = window.fetch;
                                    window.fetch = function(...args) {{
                                        const url = args[0];
                                        const options = args[1] || {{}};

                                        if (typeof url === 'string' && url.includes('output_to_document')) {{
                                            const storedData = localStorage.getItem('docFormatterFileData');
                                            if (storedData) {{
                                                try {{
                                                    const data = JSON.parse(storedData);
                                                    if (options.body) {{
                                                        const body = JSON.parse(options.body);
                                                        if (!body.file) {{
                                                            body.file = data.file;
                                                            body.file_extension = data.file_extension;
                                                            body.messages = data.messages;
                                                            body.chat_messages = data.messages;
                                                            options.body = JSON.stringify(body);
                                                            localStorage.removeItem('docFormatterFileData');
                                                            if (document.body.contains(notification)) {{
                                                                document.body.removeChild(notification);
                                                            }}
                                                        }}
                                                    }}
                                                }} catch (e) {{
                                                    console.error('Error injecting file data:', e);
                                                }}
                                            }}
                                        }}

                                        return originalFetch.apply(this, args);
                                    }};

                                    setTimeout(() => {{
                                        if (document.body.contains(notification)) {{
                                            document.body.removeChild(notification);
                                        }}
                                    }}, 5000);
                                }}
                            }}, 500);
                        }}, 500);
                    }} catch (error) {{
                        showStatus('Error: ' + error.message + '. Check browser console (F12) for details.', true);
                        console.error('Document formatter error:', error);
                    }}
                }};
                reader.readAsDataURL(file);
            }};

            document.getElementById('cancelBtn').onclick = closeModal;
            overlay.onclick = (e) => {{ if (e.target === overlay) closeModal(); }};
            overlay.onkeydown = (e) => {{ if (e.key === 'Escape') closeModal(); }};
        """

    async def action(
        self,
        body: Dict[str, Any],
        __user__: Optional[Dict[str, Any]] = None,
        __event_emitter__: Optional[Any] = None,
        __event_call__: Optional[Any] = None,
    ) -> Optional[Dict[str, Any]]:
        """
        OpenWebUI action method that:
        1. Shows a modern GUI interface when first called (button click)
        2. Accepts a document upload (DOCX or PDF)
        3. Extracts styling information from the document
        4. Formats the chat messages using the extracted styles
        5. Returns a formatted DOCX document

        Args:
            body: Dictionary containing action parameters including:
                - file: Uploaded file object (DOCX or PDF) - can be file object, file path, or base64 string
                - chat_messages: List of chat messages with 'role' and 'content' keys
                - messages: Alternative parameter name for chat messages
            __user__: Optional dictionary with user information
            __event_emitter__: Optional function to send real-time updates to the frontend
            __event_call__: Optional function for bidirectional communication (requesting user input)

        Returns:
            Dictionary with GUI HTML (to show modal), download link, or file data
        """
        # Add logging for debugging
        import logging
        logger = logging.getLogger(__name__)

        # Debug logging to stderr (visible in OpenWebUI logs)
        print(f"[DOC_FORMATTER] Action called with body keys: {list(body.keys())}", file=sys.stderr)
        print(f"[DOC_FORMATTER] Body content preview: {str(body)[:200]}", file=sys.stderr)
        try:
            logger.info(f"Action called with body keys: {list(body.keys())}")
            logger.info(f"Body content: {str(body)[:500]}")
        except:
            pass  # Don't fail if logging isn't configured

        # Extract parameters from body
        file = body.get('file')
        chat_messages = body.get('chat_messages')
        messages = body.get('messages')

        # Debug: log all body keys to see what we're receiving
        print(f"[DOC_FORMATTER] Body keys: {list(body.keys())}", file=sys.stderr)
        print(f"[DOC_FORMATTER] File in body: {file is not None}", file=sys.stderr)
        if file:
            print(f"[DOC_FORMATTER] File type: {type(file)}, length: {len(str(file)) if isinstance(file, str) else 'N/A'}", file=sys.stderr)

        # Check for file in various other possible locations
        if file is None:
            file = body.get('file_data') or body.get('uploaded_file')

        # Check if this is a function call with arguments (from JavaScript)
        if file is None and 'arguments' in body:
            try:
                import json
                args_str = body.get('arguments')
                if isinstance(args_str, str):
                    args = json.loads(args_str)
                    file = args.get('file') or args.get('file_data')
                    if file:
                        print(f"[DOC_FORMATTER] File found in function arguments", file=sys.stderr)
                        # Also get file_extension and messages from args
                        merged_kwargs['file_extension'] = args.get('file_extension', '.docx')
                        if args.get('messages'):
                            messages = args.get('messages')
            except Exception as e:
                print(f"[DOC_FORMATTER] Error parsing function arguments: {str(e)}", file=sys.stderr)

        # Merge body with all available context
        merged_kwargs = {**body}
        if __user__:
            merged_kwargs['__user__'] = __user__

        # If no file provided, use __event_call__ to prompt for file upload
        if file is None and not merged_kwargs.get('uploaded_file') and not merged_kwargs.get('file') and not merged_kwargs.get('file_data'):
            print("[DOC_FORMATTER] No file provided, prompting user for upload", file=sys.stderr)

            # Get chat messages from context
            chat_msgs = messages or chat_messages or merged_kwargs.get('messages', merged_kwargs.get('chat_messages', merged_kwargs.get('chat_history', [])))

            # Use __event_call__ with type "execute" to show modal (like prior.py)
            if __event_call__:
                print("[DOC_FORMATTER] Using __event_call__ with execute type to show file upload modal", file=sys.stderr)
                try:
                    await __event_call__({
                        "type": "execute",
                        "data": {
                            "code": self._get_file_upload_modal(chat_msgs)
                        }
                    })
                    return {"message": "File upload modal displayed. Please select a document."}
                except Exception as e:
                    print(f"[DOC_FORMATTER] Error showing modal: {str(e)}", file=sys.stderr)
                    return {
                        "content": f"Error displaying file upload: {str(e)}",
                        "success": False
                    }
            else:
                # Fallback if __event_call__ not available
                print("[DOC_FORMATTER] __event_call__ not available, returning HTML", file=sys.stderr)
                return {
                    "content": "📄 Document Style Formatter\n\nPlease upload a DOCX or PDF document to extract styling from.",
                    "success": False
                }

        # File processing section - reached if file was provided or extracted above
        print(f"[DOC_FORMATTER] Starting file processing. File is None: {file is None}", file=sys.stderr)
        try:
            # Handle file upload - support multiple formats
            file_content = None
            file_ext = None

            if file is None:
                # Try to get file from merged_kwargs
                file = merged_kwargs.get('uploaded_file', merged_kwargs.get('file', None))

            if file is None:
                # Return GUI if no file
                gui_html = generate_modern_gui()
                return {
                    "type": "html",
                    "content": gui_html,
                    "success": True,
                    "message": "Please upload a document to continue."
                }

            # Handle different file input types
            if hasattr(file, 'read'):
                # File-like object
                file_content = file.read()
                filename = getattr(file, 'filename', 'document.docx')
                file_ext = os.path.splitext(filename)[1].lower()
            elif isinstance(file, str):
                # Could be file path or base64 string
                if os.path.exists(file):
                    # File path
                    with open(file, 'rb') as f:
                        file_content = f.read()
                    file_ext = os.path.splitext(file)[1].lower()
                elif file.startswith('data:'):
                    # Data URI
                    header, encoded = file.split(',', 1)
                    file_content = base64.b64decode(encoded)
                    # Extract extension from mime type
                    if 'pdf' in header.lower():
                        file_ext = '.pdf'
                    elif 'word' in header.lower() or 'docx' in header.lower():
                        file_ext = '.docx'
                else:
                    # Try as base64 (raw base64 string)
                    try:
                        file_content = base64.b64decode(file)
                        # Try to get extension from merged_kwargs
                        file_ext = merged_kwargs.get('file_extension', '.docx')
                        if file_ext and not file_ext.startswith('.'):
                            file_ext = '.' + file_ext
                    except Exception as e:
                        return {"error": f"Invalid file format: {str(e)}. Please provide a file object, file path, or base64 string."}
            elif isinstance(file, bytes):
                file_content = file
                file_ext = merged_kwargs.get('file_extension', '.docx')
            else:
                return {"error": "Invalid file format. Please upload a DOCX or PDF file."}

            if file_content is None:
                return {"error": "Could not read file content."}

            # Validate file extension
            if file_ext not in ['.docx', '.pdf']:
                return {"error": f"Unsupported file format: {file_ext}. Please upload a DOCX or PDF file."}

            # Save uploaded file temporarily
            temp_input = tempfile.NamedTemporaryFile(delete=False, suffix=file_ext)
            temp_input.write(file_content)
            temp_input_path = temp_input.name
            temp_input.close()

            try:
                # Extract styles
                extractor = DocumentStyleExtractor()
                if file_ext == '.docx':
                    styles = extractor.extract_from_docx(temp_input_path)
                elif file_ext == '.pdf':
                    styles = extractor.extract_from_pdf(temp_input_path)

                # Get chat messages from various possible sources
                if chat_messages is None:
                    chat_messages = messages

                if chat_messages is None:
                    # Try to get from merged_kwargs with various common names
                    chat_messages = (
                        merged_kwargs.get('messages') or
                        merged_kwargs.get('chat_messages') or
                        merged_kwargs.get('chat_history') or
                        merged_kwargs.get('history') or
                        []
                    )

                # Convert messages to standard format if needed
                if chat_messages:
                    formatted_messages = []
                    for msg in chat_messages:
                        if isinstance(msg, dict):
                            # Ensure it has 'role' and 'content'
                            role = msg.get('role', msg.get('from', 'user'))
                            content = msg.get('content', msg.get('text', msg.get('message', '')))
                            formatted_messages.append({'role': role, 'content': str(content)})
                        elif isinstance(msg, str):
                            # Plain string, assume it's user content
                            formatted_messages.append({'role': 'user', 'content': msg})
                    chat_messages = formatted_messages

                if not chat_messages:
                    return {"error": "No chat messages found. Please ensure chat context is available."}

                # Apply styles to chat content
                applier = DocumentStyleApplier(styles)
                applier.add_headers()
                applier.add_chat_content(chat_messages)
                applier.add_tables()
                applier.add_footers()

                # Save output document
                output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                output_path.close()

                applier.save(output_path.name)

                # Read output file
                with open(output_path.name, 'rb') as f:
                    output_bytes = f.read()

                # Clean up temporary files
                os.unlink(temp_input_path)
                os.unlink(output_path.name)

                # Prepare file data for download
                file_base64 = base64.b64encode(output_bytes).decode('utf-8')
                filename = "formatted_chat.docx"

                # Use __event_call__ to trigger download automatically (like prior.py does)
                if __event_call__:
                    download_script = f"""
                        (function() {{
                            const base64Data = '{file_base64}';
                            const filename = '{filename}';
                            const blob = new Blob([atob(base64Data)], {{ type: 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' }});
                            const url = URL.createObjectURL(blob);
                            const a = document.createElement('a');
                            a.href = url;
                            a.download = filename;
                            document.body.appendChild(a);
                            a.click();
                            URL.revokeObjectURL(url);
                            document.body.removeChild(a);
                        }})();
                    """
                    try:
                        await __event_call__({
                            "type": "execute",
                            "data": {
                                "code": download_script
                            }
                        })
                        print("[DOC_FORMATTER] Download triggered via __event_call__", file=sys.stderr)
                    except Exception as e:
                        print(f"[DOC_FORMATTER] Error triggering download: {str(e)}", file=sys.stderr)

                # Return result - OpenWebUI typically expects file data or download URL
                return {
                    "success": True,
                    "message": "Document formatted successfully!",
                    "file": {
                        "content": file_base64,
                        "filename": filename,
                        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    },
                    "download_url": None  # Can be set if you have a file server
                }

            except Exception as e:
                # Clean up on error
                if os.path.exists(temp_input_path):
                    try:
                        os.unlink(temp_input_path)
                    except:
                        pass
                import traceback
                return {
                    "error": f"Error processing document: {str(e)}",
                    "traceback": traceback.format_exc(),
                    "success": False
                }

        except Exception as e:
            import traceback
            return {
                "error": f"Error processing document: {str(e)}",
                "traceback": traceback.format_exc(),
                "success": False
            }


# Example usage and testing
if __name__ == "__main__":
    # Test the function
    print("=" * 70)
    print("Document Style Formatter - OpenWebUI Action Function")
    print("=" * 70)
    print("\nThis function extracts styling from DOCX/PDF files and applies it to chat content.")
    print("\nSETUP:")
    print("1. Place this file (main.py) in your OpenWebUI functions directory")
    print("   - Default: ~/.open-webui/functions/")
    print("   - Or: /app/functions/ (Docker)")
    print("\n2. Install dependencies:")
    print("   pip install python-docx PyMuPDF pdf2docx")
    print("\n3. Restart OpenWebUI")
    print("\n4. The button 'Format Chat With Document Style' will appear:")
    print("   - After each chat message/response")
    print("   - In the action buttons area")
    print("\n5. Click the button to open the document upload GUI")
    print("\n" + "=" * 70)

    # Verify the Action class is properly defined
    if 'Action' in globals():
        print("\n✓ Action class found and registered")
    else:
        print("\n⚠ Action class may not be properly defined")

    print("\n" + "=" * 70)
